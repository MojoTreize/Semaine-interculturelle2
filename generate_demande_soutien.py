#!/usr/bin/env python3
"""Génère le document Word 'Demande de soutien' pour l'UGFA e.V."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Palette couleurs UGFA ─────────────────────────────────────────────────────
DARK_GREEN  = RGBColor(0x1B, 0x5E, 0x20)   # vert forêt principal
GOLD        = RGBColor(0xF9, 0xA8, 0x25)   # or accent
LIGHT_GREY  = RGBColor(0xF5, 0xF5, 0xF5)
TEXT_DARK   = RGBColor(0x21, 0x21, 0x21)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# ── Données de l'organisation ─────────────────────────────────────────────────
ORG_NAME    = "Union de la Guinée Forestière en Allemagne"
ORG_SHORT   = "UGFA e.V."
EVENT_NAME  = ("Semaine de Coopération Internationale et de\n"
               "Dialogue Interculturelle de la Guinée Forestière en Allemagne")
EVENT_DATES = "1ᵉʳ – 11 octobre 2026"
EVENT_PLACE = "Leonie-Reygers-Terrasse, 44137 Dortmund, Allemagne"
EMAIL       = "contact@ugfa-ev.org"
PHONE       = "+49 151 926 242 516"
WEBSITE     = "ugfa-ev.org"
THEME       = ("Participation des ressortissants de la Guinée Forestière en Allemagne "
               "au développement durable de leur région d'origine dans le contexte "
               "du Projet Simandou 2040 en République de Guinée.")

OUTPUT_FILE = "/home/nestor/Semaine-interculturelle2/Demande_de_Soutien_UGFA_2026.docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Remplit le fond d'une cellule de tableau."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_run_font(run, bold=False, italic=False, size_pt=11,
                 color: RGBColor = None, name="Calibri"):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = name
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text="", bold=False, italic=False, size_pt=11,
                  color: RGBColor = None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6, keep_together=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if keep_together:
        pf.keep_together = True
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic, size_pt=size_pt, color=color)
    return p


def add_horizontal_rule(doc, color_hex="1B5E20", thickness=12):
    """Ajoute une fine ligne de séparation."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def build_header(section):
    """Construit l'en-tête avec bandeau couleur et infos événement."""
    header = section.header
    header.is_linked_to_previous = False

    # Vider le contenu existant
    for p in header.paragraphs:
        p.clear()

    # Tableau 1 colonne = bandeau vert plein fond
    tbl = header.add_table(rows=1, cols=2, width=Inches(7))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Colonne gauche : nom + événement
    left = tbl.cell(0, 0)
    left.width = Inches(5.2)
    set_cell_bg(left, "1B5E20")
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = left.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(2)
    p1.paragraph_format.left_indent  = Cm(0.4)
    r1 = p1.add_run(ORG_NAME.upper())
    set_run_font(r1, bold=True, size_pt=10, color=WHITE, name="Calibri")

    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    p2.paragraph_format.left_indent  = Cm(0.4)
    r2 = p2.add_run(EVENT_NAME.replace("\n", " — "))
    set_run_font(r2, italic=True, size_pt=8, color=GOLD, name="Calibri")

    # Colonne droite : dates + lieu
    right = tbl.cell(0, 1)
    right.width = Inches(1.8)
    set_cell_bg(right, "F9A825")
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p3 = right.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(6)
    p3.paragraph_format.space_after  = Pt(2)
    r3 = p3.add_run("Dortmund 2026")
    set_run_font(r3, bold=True, size_pt=9, color=DARK_GREEN, name="Calibri")

    p4 = right.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after  = Pt(6)
    r4 = p4.add_run(EVENT_DATES)
    set_run_font(r4, size_pt=7.5, color=DARK_GREEN, name="Calibri")

    # Supprimer les bordures du tableau d'en-tête
    for row in tbl.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "none")
                el.set(qn("w:sz"),    "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBdr.append(el)
            tcPr.append(tcBdr)


def build_footer(section):
    """Construit le pied de page avec contacts, site web et numéro de page."""
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    # Ligne de séparation
    pSep = footer.paragraphs[0]
    pPr  = pSep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "12")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "1B5E20")
    pBdr.append(top)
    pPr.append(pBdr)
    pSep.paragraph_format.space_after = Pt(3)

    # Tableau 3 colonnes : coordonnées | événement centré | page
    tbl = footer.add_table(rows=1, cols=3, width=Inches(7))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(2.8), Inches(2.4), Inches(1.8)]
    cells  = tbl.rows[0].cells
    for i, w in enumerate(widths):
        cells[i].width = w

    # Colonne gauche – coordonnées
    cl = cells[0]
    cl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pL1 = cl.paragraphs[0]
    pL1.paragraph_format.space_before = Pt(2)
    r = pL1.add_run(f"✉  {EMAIL}   ☎  {PHONE}")
    set_run_font(r, size_pt=7.5, color=DARK_GREEN)

    pL2 = cl.add_paragraph()
    pL2.paragraph_format.space_before = Pt(0)
    pL2.paragraph_format.space_after  = Pt(2)
    r2 = pL2.add_run(f"🌐  {WEBSITE}   📍  Dortmund, Allemagne")
    set_run_font(r2, size_pt=7.5, color=TEXT_DARK)

    # Colonne centrale – nom org
    cm = cells[1]
    cm.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pC = cm.paragraphs[0]
    pC.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pC.paragraph_format.space_before = Pt(2)
    pC.paragraph_format.space_after  = Pt(2)
    r3 = pC.add_run(f"{ORG_SHORT} — Dortmund 2026")
    set_run_font(r3, bold=True, size_pt=7.5, color=DARK_GREEN)

    # Colonne droite – numéro de page
    cr = cells[2]
    cr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pR = cr.paragraphs[0]
    pR.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pR.paragraph_format.space_before = Pt(2)
    pR.paragraph_format.space_after  = Pt(2)
    r4 = pR.add_run("Page ")
    set_run_font(r4, size_pt=7.5, color=TEXT_DARK)
    # Champ PAGE
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz  = OxmlElement("w:sz"); sz.set(qn("w:val"), "15")
    rPr.append(sz)
    new_run.append(rPr)
    new_run.append(fldChar1)
    new_run.append(instrText)
    new_run.append(fldChar2)
    pR._p.append(new_run)

    # Supprimer les bordures du tableau de pied de page
    for row in tbl.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "none")
                el.set(qn("w:sz"),    "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBdr.append(el)
            tcPr.append(tcBdr)


def add_info_box(doc, title, lines):
    """Insère un encadré coloré (tableau 1 cellule)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "E8F5E9")

    p_title = cell.paragraphs[0]
    p_title.paragraph_format.left_indent   = Cm(0.3)
    p_title.paragraph_format.space_before  = Pt(4)
    p_title.paragraph_format.space_after   = Pt(2)
    rt = p_title.add_run(title)
    set_run_font(rt, bold=True, size_pt=10, color=DARK_GREEN)

    for line in lines:
        pl = cell.add_paragraph()
        pl.paragraph_format.left_indent  = Cm(0.5)
        pl.paragraph_format.space_before = Pt(1)
        pl.paragraph_format.space_after  = Pt(1)
        rl = pl.add_run(line)
        set_run_font(rl, size_pt=10, color=TEXT_DARK)

    cell.add_paragraph().paragraph_format.space_after = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_sponsorship_table(doc):
    """Tableau des niveaux de sponsoring."""
    headers = ["Niveau", "Montant", "Avantages principaux"]
    rows_data = [
        ["Bronze",     "100 €",         "Mention sur les supports numériques"],
        ["Silver",     "500 €",         "Logo + mention + 2 invitations officielles"],
        ["Gold",       "1 000 €",       "Logo premium + espace d'exposition + 4 invitations"],
        ["Stratégique","2 000 € et +",  "Partenariat institutionnel complet + visibilité maximale"],
    ]
    row_colors = ["FFF9C4", "FFF9C4", "FFF9C4", "FFF9C4"]
    header_color = "1B5E20"

    tbl = doc.add_table(rows=1 + len(rows_data), cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-tête
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_color)
        ph = hdr_cells[i].paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_before = Pt(4)
        ph.paragraph_format.space_after  = Pt(4)
        rh = ph.add_run(h)
        set_run_font(rh, bold=True, size_pt=10, color=WHITE)

    # Lignes
    for ri, (row_data, bg) in enumerate(zip(rows_data, row_colors)):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            set_cell_bg(cells[ci], bg if ri % 2 == 0 else "FFFFFF")
            pc = cells[ci].paragraphs[0]
            pc.paragraph_format.space_before = Pt(3)
            pc.paragraph_format.space_after  = Pt(3)
            pc.paragraph_format.left_indent  = Cm(0.2)
            rc = pc.add_run(val)
            bold = (ci == 0)
            color = DARK_GREEN if ci == 0 else TEXT_DARK
            set_run_font(rc, bold=bold, size_pt=10, color=color)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ── Construction du document ──────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Marges de page
    section = doc.sections[0]
    section.top_margin    = Cm(3.5)
    section.bottom_margin = Cm(2.8)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

    # En-tête & pied de page
    build_header(section)
    build_footer(section)

    # ── Objet ─────────────────────────────────────────────────────────────────
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ref.paragraph_format.space_after = Pt(4)
    r_ref = p_ref.add_run("Dortmund, le ___________________________")
    set_run_font(r_ref, size_pt=10, color=TEXT_DARK)

    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.space_before = Pt(6)
    p_obj.paragraph_format.space_after  = Pt(12)
    p_obj.add_run("Objet : ").font.bold = True
    r_obj = p_obj.add_run(
        "Demande de soutien — Semaine de Coopération Internationale "
        "et de Dialogue Interculturelle de la Guinée Forestière en Allemagne, Dortmund 2026"
    )
    set_run_font(r_obj, bold=True, size_pt=11, color=DARK_GREEN)

    add_horizontal_rule(doc)

    # ── Bloc destinataire ─────────────────────────────────────────────────────
    add_paragraph(doc, "À l'attention de :", bold=True, size_pt=11,
                  color=TEXT_DARK, space_before=10, space_after=2)
    add_paragraph(doc, "M. / Mme ________________________________________",
                  size_pt=11, space_after=1)
    add_paragraph(doc, "Organisation / Institution : _____________________________",
                  size_pt=11, space_after=1)
    add_paragraph(doc, "Adresse : ___________________________________________",
                  size_pt=11, space_after=10)

    # ── Formule d'appel ───────────────────────────────────────────────────────
    add_paragraph(doc, "Madame, Monsieur,", bold=False, size_pt=11,
                  space_before=4, space_after=10)

    # ── Paragraphe d'introduction ─────────────────────────────────────────────
    add_paragraph(
        doc,
        "L'Union de la Guinée Forestière en Allemagne (UGFA e.V.) a l'honneur de vous "
        "présenter ce dossier de demande de soutien dans le cadre de l'organisation de la "
        "Semaine de Coopération Internationale et de Dialogue Interculturelle de la Guinée "
        "Forestière en Allemagne, prévue du 1er au 11 octobre 2026 à Dortmund, Allemagne.",
        size_pt=11, space_after=8
    )

    add_paragraph(
        doc,
        "Cet événement de dimension internationale réunira la diaspora guinéenne d'Allemagne, "
        "des représentants institutionnels, des experts sectoriels et des partenaires "
        "stratégiques autour d'une thématique centrale :",
        size_pt=11, space_after=4
    )

    p_theme = doc.add_paragraph()
    p_theme.paragraph_format.left_indent   = Cm(1.0)
    p_theme.paragraph_format.right_indent  = Cm(1.0)
    p_theme.paragraph_format.space_before  = Pt(2)
    p_theme.paragraph_format.space_after   = Pt(10)
    r_theme = p_theme.add_run(f"« {THEME} »")
    set_run_font(r_theme, italic=True, size_pt=10.5, color=DARK_GREEN)

    # ── Encadré infos événement ───────────────────────────────────────────────
    add_info_box(doc, "Informations clés de l'événement", [
        f"📅  Dates         :  {EVENT_DATES}",
        f"📍  Lieu           :  {EVENT_PLACE}",
        f"🎯  Thème         :  Projet Simandou 2040 & développement durable",
        f"👥  Public visé  :  Diaspora, institutions, experts, partenaires internationaux",
    ])

    # ── Objectifs ─────────────────────────────────────────────────────────────
    add_paragraph(doc, "Objectifs de l'événement", bold=True, size_pt=12,
                  color=DARK_GREEN, space_before=4, space_after=4)
    add_horizontal_rule(doc, color_hex="F9A825", thickness=6)

    objectives = [
        "Mobiliser les ressortissants de la diaspora autour de projets concrets pour la Guinée Forestière.",
        "Favoriser des partenariats durables entre acteurs allemands et guinéens.",
        "Promouvoir la transparence, l'inclusion et la durabilité dans les projets de développement.",
        "Structurer une feuille de route d'investissement territorial 2026–2030.",
        "Contribuer à l'intégration des retombées du Projet Simandou 2040 au bénéfice des populations locales.",
    ]
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(obj)
        set_run_font(r, size_pt=10.5, color=TEXT_DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Objet de la demande ───────────────────────────────────────────────────
    add_paragraph(doc, "Nature de la demande de soutien", bold=True, size_pt=12,
                  color=DARK_GREEN, space_before=6, space_after=4)
    add_horizontal_rule(doc, color_hex="F9A825", thickness=6)

    add_paragraph(
        doc,
        "Dans le cadre de la préparation de cet événement, nous sollicitons votre soutien "
        "selon l'une ou plusieurs des modalités suivantes :",
        size_pt=11, space_after=6
    )

    modalities = [
        ("Soutien financier",
         "contribution directe au budget de l'événement selon les niveaux de partenariat définis."),
        ("Soutien en nature",
         "mise à disposition de locaux, équipements, services logistiques ou prestations."),
        ("Soutien institutionnel",
         "caution officielle, lettre de recommandation ou appui à la communication."),
        ("Partenariat médiatique",
         "relais de l'événement sur vos canaux de communication et plateformes."),
        ("Partenariat académique / expertise",
         "participation d'intervenants, contribution à la conception du programme."),
    ]

    for title, desc in modalities:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        r_bull = p.add_run("▶  ")
        set_run_font(r_bull, size_pt=10, color=GOLD)
        r_t = p.add_run(f"{title} : ")
        set_run_font(r_t, bold=True, size_pt=10.5, color=DARK_GREEN)
        r_d = p.add_run(desc)
        set_run_font(r_d, size_pt=10.5, color=TEXT_DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── Niveaux de sponsoring ─────────────────────────────────────────────────
    add_paragraph(doc, "Niveaux de sponsoring financier", bold=True, size_pt=12,
                  color=DARK_GREEN, space_before=4, space_after=6)
    add_sponsorship_table(doc)

    # ── Contreparties & visibilité ────────────────────────────────────────────
    add_paragraph(doc, "Contreparties proposées à nos partenaires", bold=True, size_pt=12,
                  color=DARK_GREEN, space_before=4, space_after=4)
    add_horizontal_rule(doc, color_hex="F9A825", thickness=6)

    benefits = [
        "Visibilité institutionnelle sur tous les supports officiels (flyers, bannières, site web, réseaux sociaux).",
        "Mention dans le discours d'ouverture et les communications officielles de l'événement.",
        "Accès à un réseau qualifié : diaspora, décideurs, experts et acteurs économiques.",
        "Invitation officielle aux sessions plénières et cérémonies de l'événement.",
        "Remise d'un certificat de partenariat signé par le comité d'organisation.",
        "Rapport d'impact post-événement transmis à chaque partenaire.",
    ]
    for b in benefits:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(b)
        set_run_font(r, size_pt=10.5, color=TEXT_DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── Conclusion ────────────────────────────────────────────────────────────
    add_paragraph(
        doc,
        "Nous demeurons disponibles pour vous rencontrer à votre convenance afin de présenter "
        "ce projet plus en détail et de convenir ensemble des modalités de partenariat. "
        "Dans l'attente de votre réponse favorable, nous vous prions d'agréer, Madame, Monsieur, "
        "l'expression de nos sentiments distingués.",
        size_pt=11, space_before=6, space_after=16
    )

    # ── Signature ─────────────────────────────────────────────────────────────
    tbl_sig = doc.add_table(rows=1, cols=2)
    tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells_sig = tbl_sig.rows[0].cells
    cells_sig[0].width = Inches(3.5)
    cells_sig[1].width = Inches(3.5)

    for ci, (label, sub) in enumerate([
        ("Le Président / La Présidente", "Union de la Guinée Forestière en Allemagne"),
        ("Le Secrétaire Général / La Secrétaire Générale", ORG_SHORT),
    ]):
        pc = cells_sig[ci].paragraphs[0]
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(2)
        rc = pc.add_run(label)
        set_run_font(rc, bold=True, size_pt=10, color=DARK_GREEN)

        p2c = cells_sig[ci].add_paragraph()
        p2c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2c.paragraph_format.space_after = Pt(2)
        r2c = p2c.add_run(sub)
        set_run_font(r2c, italic=True, size_pt=9, color=TEXT_DARK)

        p3c = cells_sig[ci].add_paragraph()
        p3c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3c.paragraph_format.space_before = Pt(16)
        p3c.paragraph_format.space_after  = Pt(2)
        r3c = p3c.add_run("Signature & cachet : _______________________")
        set_run_font(r3c, size_pt=9.5, color=TEXT_DARK)

    # Supprimer bordures du tableau signature
    for row in tbl_sig.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "none")
                el.set(qn("w:sz"),    "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBdr.append(el)
            tcPr.append(tcBdr)

    # ── Pied de document ──────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_horizontal_rule(doc)
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note.paragraph_format.space_before = Pt(4)
    r_note = p_note.add_run(
        f"Pour tout renseignement complémentaire : {EMAIL}  |  {PHONE}  |  {WEBSITE}"
    )
    set_run_font(r_note, size_pt=8.5, color=DARK_GREEN)

    doc.save(OUTPUT_FILE)
    print(f"✅  Document généré : {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
