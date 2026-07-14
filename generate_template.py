#!/usr/bin/env python3
"""Template Word vierge – en-tête logo + infos asso, pied de page contacts."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_GREEN = RGBColor(0x1B, 0x5E, 0x20)
GOLD       = RGBColor(0xF9, 0xA8, 0x25)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY       = RGBColor(0x60, 0x60, 0x60)
TEXT_DARK  = RGBColor(0x21, 0x21, 0x21)

LOGO_PATH  = "/home/nestor/Semaine-interculturelle2/assets/images/logo.jpeg"
OUTPUT     = "/home/nestor/Semaine-interculturelle2/Template_UGFA_2026.docx"

# Associations membres (nom, numéro d'enregistrement)
MEMBRES = [
    ("UGS – Union de la Guinée Forestière e.V.", "VR 36787 B"),
    ("CRGSA e.V.",                               "VR 4896"),
    ("Lelona e.V.",                              "VR 41072 B"),
    ("Alga e.V.",                                "VR 703617"),
    ("Zaly e.V.",                                "VR 5750"),
]

ORG_FULL   = "Union de la Guinée Forestière en Allemagne"
ORG_SHORT  = "UGFA e.V."
EVENT_NAME = "Semaine de Coopération Internationale et de Dialogue Interculturelle"
EVENT_SUB  = "de la Guinée Forestière en Allemagne — Dortmund 2026"
EVENT_DATES= "1er – 11 octobre 2026"
EVENT_PLACE= "Leonie-Reygers-Terrasse, 44137 Dortmund, Allemagne"
EMAIL      = "contact@ugfa-ev.org"
PHONE      = "+49 151 926 242 516"
WEBSITE    = "ugfa-ev.org"


def no_borders(tbl):
    """Supprime toutes les bordures d'un tableau."""
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ("top","left","bottom","right","insideH","insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "none")
                el.set(qn("w:sz"),    "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBdr.append(el)
            tcPr.append(tcBdr)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def run(para, text, bold=False, italic=False, size=11,
        color=None, font="Calibri"):
    r = para.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.name   = font
    r.font.size   = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def hr(container, color_hex="1B5E20", thick=12, before=2, after=2):
    p  = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thick))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def build_header(section):
    hdr = section.header
    hdr.is_linked_to_previous = False
    for p in hdr.paragraphs:
        p.clear()

    # Tableau 2 colonnes : logo | texte asso
    tbl = hdr.add_table(rows=1, cols=2, width=Inches(7.0))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(tbl)

    col_logo = tbl.cell(0, 0)
    col_logo.width = Inches(1.4)
    col_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    col_text = tbl.cell(0, 1)
    col_text.width = Inches(5.6)
    col_text.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ── Logo ──────────────────────────────────────────────────────────────────
    p_logo = col_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(2)
    p_logo.paragraph_format.space_after  = Pt(2)
    try:
        r_logo = p_logo.add_run()
        r_logo.add_picture(LOGO_PATH, width=Cm(2.6))
    except Exception as e:
        run(p_logo, "[Logo]", bold=True, size=9, color=DARK_GREEN)

    # ── Texte asso ────────────────────────────────────────────────────────────
    p1 = col_text.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(1)
    p1.paragraph_format.left_indent  = Cm(0.3)
    run(p1, ORG_FULL.upper(), bold=True, size=11, color=DARK_GREEN)

    p2 = col_text.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(1)
    p2.paragraph_format.left_indent  = Cm(0.3)
    run(p2, EVENT_NAME, italic=True, size=9.5, color=GREY)

    p3 = col_text.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(4)
    p3.paragraph_format.left_indent  = Cm(0.3)
    run(p3, EVENT_SUB, italic=True, size=9.5, color=GREY)

    # ── Bandeau doré sous le tableau ─────────────────────────────────────────
    tbl2 = hdr.add_table(rows=1, cols=1, width=Inches(7.0))
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(tbl2)
    cell2 = tbl2.cell(0, 0)
    set_cell_bg(cell2, "1B5E20")

    p_band = cell2.paragraphs[0]
    p_band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_band.paragraph_format.space_before = Pt(3)
    p_band.paragraph_format.space_after  = Pt(3)
    run(p_band, f"  {EVENT_DATES}  ·  {EVENT_PLACE}  ",
        size=8.5, color=GOLD, bold=False)


def build_footer(section):
    ftr = section.footer
    ftr.is_linked_to_previous = False
    for p in ftr.paragraphs:
        p.clear()

    # ── Ligne de séparation verte ─────────────────────────────────────────────
    p_sep = ftr.paragraphs[0]
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after  = Pt(2)
    pPr  = p_sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "12")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "1B5E20")
    pBdr.append(top)
    pPr.append(pBdr)

    # ── Ligne 1 : associations membres ───────────────────────────────────────
    p_membres = ftr.add_paragraph()
    p_membres.paragraph_format.space_before = Pt(2)
    p_membres.paragraph_format.space_after  = Pt(3)
    p_membres.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run(p_membres, "Associations membres : ", bold=True, size=7.5, color=DARK_GREEN)
    for i, (nom, vr) in enumerate(MEMBRES):
        if i > 0:
            run(p_membres, "  |  ", size=7.5, color=GREY)
        run(p_membres, nom, bold=True, size=7.5, color=TEXT_DARK)
        run(p_membres, f" ({vr})", size=7, color=GREY)

    # ── Ligne séparation fine ─────────────────────────────────────────────────
    p_mid = ftr.add_paragraph()
    p_mid.paragraph_format.space_before = Pt(0)
    p_mid.paragraph_format.space_after  = Pt(1)
    pPr2  = p_mid._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    bot2  = OxmlElement("w:bottom")
    bot2.set(qn("w:val"),   "single")
    bot2.set(qn("w:sz"),    "4")
    bot2.set(qn("w:space"), "1")
    bot2.set(qn("w:color"), "CCCCCC")
    pBdr2.append(bot2)
    pPr2.append(pBdr2)

    # ── Ligne 2 : contacts | org | page ──────────────────────────────────────
    tbl = ftr.add_table(rows=1, cols=3, width=Inches(7.0))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(tbl)

    cells = tbl.rows[0].cells
    cells[0].width = Inches(3.2)
    cells[1].width = Inches(2.0)
    cells[2].width = Inches(1.8)

    for c in cells:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Gauche – contacts
    pL = cells[0].paragraphs[0]
    pL.paragraph_format.space_before = Pt(1)
    pL.paragraph_format.space_after  = Pt(0)
    run(pL, f"{EMAIL}  ·  {PHONE}", size=7.5, color=GREY)
    pL2 = cells[0].add_paragraph()
    pL2.paragraph_format.space_before = Pt(0)
    pL2.paragraph_format.space_after  = Pt(2)
    run(pL2, f"{WEBSITE}  ·  Dortmund, Allemagne", size=7.5, color=GREY)

    # Centre – nom org
    pC = cells[1].paragraphs[0]
    pC.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pC.paragraph_format.space_before = Pt(1)
    pC.paragraph_format.space_after  = Pt(2)
    run(pC, ORG_SHORT, bold=True, size=8, color=DARK_GREEN)

    # Droite – numéro de page
    pR = cells[2].paragraphs[0]
    pR.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pR.paragraph_format.space_before = Pt(1)
    pR.paragraph_format.space_after  = Pt(2)
    run(pR, "Page ", size=7.5, color=GREY)
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr    = OxmlElement("w:instrText"); instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    nr = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz  = OxmlElement("w:sz"); sz.set(qn("w:val"), "15")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "606060")
    rPr.append(sz); rPr.append(col)
    nr.append(rPr); nr.append(fldChar1); nr.append(instr); nr.append(fldChar2)
    pR._p.append(nr)


def build():
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin    = Cm(4.0)
    sec.bottom_margin = Cm(4.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)

    build_header(sec)
    build_footer(sec)

    # Corps vide – juste un paragraphe pour que Word ouvre le fichier proprement
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)

    doc.save(OUTPUT)
    print(f"✅  Template généré : {OUTPUT}")


if __name__ == "__main__":
    build()
